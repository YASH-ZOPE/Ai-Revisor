"""
DOCX Reader Utility.
Extracts clean text from Microsoft Word documents using python-docx.
"""
from typing import BinaryIO
import docx

def extract_text_from_docx(file_obj: BinaryIO) -> str:
    """
    Extracts text from a DOCX file object.
    
    Args:
        file_obj: A binary file-like object (e.g., BytesIO from Streamlit file uploader).
        
    Returns:
        str: Extracted and cleaned text.
        
    Raises:
        ValueError: If the file is empty or text extraction yields no content.
    """
    try:
        # Reset stream position just in case
        file_obj.seek(0)
        doc = docx.Document(file_obj)
        
        extracted_text = []
        
        # Process document elements in order if possible, or extract paragraphs and tables
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                extracted_text.append(text)
                
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    # Strip cells to clean up whitespace
                    cell_text = cell.text.strip()
                    if cell_text:
                        # Avoid duplicates if cells are merged
                        if not row_text or row_text[-1] != cell_text:
                            row_text.append(cell_text)
                if row_text:
                    extracted_text.append(" | ".join(row_text))
                    
        full_text = "\n\n".join(extracted_text).strip()
        if not full_text:
            raise ValueError("The Word document contains no readable text.")
            
        return full_text
    except Exception as e:
        if isinstance(e, ValueError):
            raise e
        raise ValueError(f"Error reading DOCX file: {str(e)}")
